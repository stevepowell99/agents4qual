"""
Convert interview DOCX files in /sources to one Markdown file per source in /sources_md.

Kept intentionally simple:
- Extracts paragraphs in order (no styling).
- Writes a short metadata header derived from the filename.
- Treats each DOCX as an independent source (no concatenation).
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCES_DIR = ROOT / "sources"
OUT_DIR = ROOT / "sources_md"


@dataclass(frozen=True)
class ParsedName:
    raw_filename: str
    interview_number: int | None
    gender: str | None
    age: int | None
    borough: str | None
    parts_hint: str | None


def _clean(s: str) -> str:
    # Fix common mojibake seen in some filenames (best-effort, minimal).
    return s.replace("ÔÇô", "–").replace("ÔÇô", "-").strip()


def parse_filename(filename: str) -> ParsedName:
    s = _clean(filename)

    # Interview number patterns.
    m_num = re.search(r"\bInterview\s*(\d+)\b", s, flags=re.IGNORECASE)
    if not m_num:
        m_num = re.search(r"\bParticipant\s*(\d+)\b", s, flags=re.IGNORECASE)
    if not m_num:
        m_num = re.search(r"\bParticipant\s*number\s*(\d+)\b", s, flags=re.IGNORECASE)
    interview_number = int(m_num.group(1)) if m_num else None

    # Gender patterns.
    gender = None
    if re.search(r"\bFemale\b", s, flags=re.IGNORECASE):
        gender = "Female"
    elif re.search(r"\bMale\b", s, flags=re.IGNORECASE):
        gender = "Male"

    # Age patterns.
    # Filenames typically contain "... Male, 24, Newham ..." — we prefer that comma-delimited slot.
    age = None
    m_age = re.search(r",\s*(\d{2})\s*,", s)
    if m_age:
        age = int(m_age.group(1))
    else:
        # Fallback: after gender token if present.
        m_age2 = re.search(r"\b(?:Male|Female)\b\s*,\s*(\d{2})\b", s, flags=re.IGNORECASE)
        if m_age2:
            age = int(m_age2.group(1))

    # Borough patterns (from filenames in this repo).
    borough = None
    if re.search(r"\bNewham\b", s, flags=re.IGNORECASE):
        borough = "Newham"
    elif re.search(r"\bHackney\b", s, flags=re.IGNORECASE):
        borough = "Hackney"
    elif re.search(r"\bTower\s*Hamlets\b", s, flags=re.IGNORECASE):
        borough = "Tower Hamlets"
    elif re.search(r"\bB&D\b|\bB\s*&\s*D\b|Barking\s*&\s*Dagenham", s, flags=re.IGNORECASE):
        borough = "Barking & Dagenham"

    # Part hints.
    parts_hint = None
    if re.search(r"part\s*1\s*&\s*2|parts?\s*1\s*&\s*2|\(part\s*1\s*&\s*2\)|\(parts?\s*1\s*&\s*2\)", s, flags=re.IGNORECASE):
        parts_hint = "Part 1 & 2"
    elif re.search(r"\bpart\s*1\b", s, flags=re.IGNORECASE):
        parts_hint = "Part 1"
    elif re.search(r"\bpart\s*2\b", s, flags=re.IGNORECASE):
        parts_hint = "Part 2"

    return ParsedName(
        raw_filename=filename,
        interview_number=interview_number,
        gender=gender,
        age=age,
        borough=borough,
        parts_hint=parts_hint,
    )


def docx_to_paragraphs(path: Path) -> list[str]:
    doc = Document(str(path))
    paras: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            paras.append(text)
    # Some DOCX store the transcript inside tables; include those too.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = (p.text or "").strip()
                    if text:
                        paras.append(text)
    # Minimal de-duplication for adjacent repeats introduced by table structures.
    deduped: list[str] = []
    for t in paras:
        if not deduped or deduped[-1] != t:
            deduped.append(t)
    return deduped


def make_out_name(p: ParsedName, fallback_stem: str) -> str:
    """
    Make a base output name.

    IMPORTANT: The corpus contains multiple DOCX files that can share the same interview_number
    (e.g. a separate "part 2" file). So the caller must ensure uniqueness (e.g. by de-duping).
    """
    safe_stem = re.sub(r"[^A-Za-z0-9_]+", "_", fallback_stem).strip("_")
    if p.interview_number is None:
        return f"Interview_{safe_stem}.md"

    # Prefer a structured suffix when we can infer the part(s), otherwise use the stem.
    suffix = None
    if p.parts_hint:
        suffix = (
            p.parts_hint.lower()
            .replace("&", "and")
            .replace(" ", "")
        )  # e.g. "part1and2", "part2"
    else:
        suffix = safe_stem.lower()

    return f"Interview_{p.interview_number:02d}_{suffix}.md"


def render_md(p: ParsedName, paragraphs: Iterable[str]) -> str:
    meta_lines = [
        "---",
        f"source_filename: {p.raw_filename}",
        f"interview_number: {p.interview_number if p.interview_number is not None else ''}",
        f"gender: {p.gender or ''}",
        f"age: {p.age if p.age is not None else ''}",
        f"borough: {p.borough or ''}",
        f"parts_hint: {p.parts_hint or ''}",
        "---",
        "",
        "# Transcript",
        "",
    ]
    body = "\n\n".join(paragraphs).strip() + "\n"
    return "\n".join(meta_lines) + body


def main() -> int:
    if not SOURCES_DIR.exists():
        raise SystemExit(f"Missing sources dir: {SOURCES_DIR}")

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    docx_files = sorted([p for p in SOURCES_DIR.glob("*.docx") if p.is_file()])
    if not docx_files:
        raise SystemExit(f"No .docx files found in {SOURCES_DIR}")

    used_names: dict[str, int] = {}
    mapping_rows: list[tuple[str, str]] = []

    for docx_path in docx_files:
        parsed = parse_filename(docx_path.name)
        out_name = make_out_name(parsed, docx_path.stem)

        # Ensure uniqueness within this run (handle truly duplicated stems/parts).
        if out_name in used_names:
            used_names[out_name] += 1
            stem = out_name[:-3]  # strip ".md"
            out_name = f"{stem}_dup{used_names[out_name]}.md"
        else:
            used_names[out_name] = 1

        out_path = OUT_DIR / out_name

        paragraphs = docx_to_paragraphs(docx_path)
        out_md = render_md(parsed, paragraphs)
        out_path.write_text(out_md, encoding="utf-8")
        mapping_rows.append((docx_path.name, out_name))

    # Write a simple mapping file for transparency/debugging.
    mapping_path = OUT_DIR / "_conversion_map.md"
    lines = [
        "# DOCX → Markdown conversion map",
        "",
        "This file is generated by `scripts/convert_sources_docx_to_md.py`.",
        "",
        "| source_docx | output_md |",
        "|---|---|",
    ]
    for src, out in mapping_rows:
        lines.append(f"| {src} | {out} |")
    mapping_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

